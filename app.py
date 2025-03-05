import streamlit as st
import re
import io
from googleapiclient.discovery import build
import docx

# Get the API key from Streamlit secrets
API_KEY = st.secrets["YOUTUBE_API_KEY"]

def extract_video_id(url):
    """
    Extracts the video ID from a YouTube URL.
    """
    regex = r"(?:v=|\/)([0-9A-Za-z_-]{11}).*"
    match = re.search(regex, url)
    if match:
        return match.group(1)
    return None

def get_comments(video_id, api_key):
    """
    Uses the YouTube Data API to retrieve all top-level comments for the given video.
    """
    comments = []
    youtube = build('youtube', 'v3', developerKey=api_key)
    nextPageToken = None

    while True:
        request = youtube.commentThreads().list(
            part="snippet",
            videoId=video_id,
            maxResults=100,  # maximum allowed per request
            pageToken=nextPageToken
        )
        response = request.execute()

        for item in response.get("items", []):
            snippet = item["snippet"]["topLevelComment"]["snippet"]
            comment_text = snippet.get("textDisplay")
            like_count = snippet.get("likeCount")
            comments.append({
                "text": comment_text,
                "like_count": like_count
            })

        nextPageToken = response.get("nextPageToken")
        if not nextPageToken:
            break

    return comments

def create_docx(comments):
    """
    Creates a DOCX file in memory containing the comments.
    """
    doc = docx.Document()
    for comment in comments:
        doc.add_paragraph(f"Likes: {comment['like_count']}")
        doc.add_paragraph(comment["text"])
        doc.add_paragraph("-" * 40)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

st.title("YouTube Comments Downloader")

# Get YouTube video URL from the user
youtube_url = st.text_input("Enter YouTube video URL:")

if youtube_url:
    video_id = extract_video_id(youtube_url)
    if not video_id:
        st.error("Invalid YouTube URL. Please enter a valid URL.")
    else:
        st.info("Fetching comments... This may take a moment.")
        try:
            comments = get_comments(video_id, API_KEY)
            st.write(f"Retrieved {len(comments)} comments.")

            if comments:
                # Sort comments by like count in descending order
                sorted_comments = sorted(comments, key=lambda x: x["like_count"], reverse=True)
                docx_data = create_docx(sorted_comments)
                
                st.download_button(
                    label="Download Comments as DOCX",
                    data=docx_data,
                    file_name="comments.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.warning("No comments found for this video.")
        except Exception as e:
            st.error(f"An error occurred: {str(e)}")
